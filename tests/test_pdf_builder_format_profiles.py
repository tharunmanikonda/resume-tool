from pathlib import Path

from docx import Document
from docx.shared import Inches

from pdf_builder import build_resume_docx


def assert_near(actual, expected, tolerance=200):
    assert abs(int(actual) - int(expected)) <= tolerance


def sample_resume():
    return {
        "name": "Test Candidate",
        "title": "Software Engineer",
        "contact": {
            "location": "Austin, TX",
            "phone": "555-0100",
            "email": "candidate@example.com",
        },
        "summary": "Builds reliable systems.",
        "technical_skills": [
            {"category": "Languages", "items": "Python, JavaScript"},
        ],
        "experience": [
            {
                "company": "Example",
                "location": "Remote",
                "title": "Engineer",
                "dates": "2024 - Present",
                "bullets": ["Built automation."],
            },
        ],
        "projects": [
            {"name": "Resume Tool", "bullets": ["Generated resumes."]},
        ],
        "education": [
            {
                "degree": "MS Computer Science",
                "institution": "Example University",
                "dates": "2023",
            },
        ],
        "certifications": ["AWS Certified"],
    }


def test_gmail_format_profile_applies_docx_layout(tmp_path):
    output_docx = Path(tmp_path) / "gmail.docx"

    build_resume_docx(sample_resume(), str(output_docx), format_profile="gmail")

    doc = Document(str(output_docx))
    section = doc.sections[0]

    assert_near(section.top_margin, Inches(0.25))
    assert_near(section.bottom_margin, Inches(0.25))
    assert_near(section.left_margin, Inches(0.28))
    assert_near(section.right_margin, Inches(0.28))
    assert doc.styles["Normal"].font.name == "Times New Roman"


def test_skill_item_arrays_render_as_comma_separated_text(tmp_path):
    output_docx = Path(tmp_path) / "skills.docx"
    resume = sample_resume()
    resume["technical_skills"] = [
        {"category": "Programming Languages", "items": ["TypeScript", "JavaScript", "Java", "SQL"]},
    ]

    build_resume_docx(resume, str(output_docx))

    doc = Document(str(output_docx))
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "Programming Languages: TypeScript, JavaScript, Java, SQL" in text
    assert "['TypeScript'" not in text
